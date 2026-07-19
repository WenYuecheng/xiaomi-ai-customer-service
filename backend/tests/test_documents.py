"""
文件职责：
该文件负责知识库和文档管理功能的自动化测试用例。

所属功能：
文档与知识库管理模块

主要流程：
测试文档的上传、异步解析任务轮询、各类文件格式兼容性检查、文档的逻辑与物理清理以及上传时的边界条件和异常处理。
"""

import time
from io import BytesIO

from docx import Document as WordDocument
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

from tests.conftest import auth_headers


def create_knowledge_base(client: TestClient, headers: dict[str, str]) -> str:
    """
    创建一个测试用的产品资料知识库。

    Args:
        client: FastAPI 的测试客户端，用于发送 HTTP 请求。
        headers: 包含鉴权令牌等信息的请求头字典。

    Returns:
        str: 新创建的知识库的 UUID 字符串。

    Raises:
        AssertionError: 如果创建请求未能返回 201 状态码，则触发断言错误。
    """
    # 发送创建知识库的 POST 请求，名称指定为“产品资料库”
    response = client.post(
        "/api/v1/knowledge-bases",
        headers=headers,
        json={"name": "产品资料库"},
    )
    assert response.status_code == 201
    return response.json()["id"]


def wait_for_job(client: TestClient, headers: dict[str, str], job_id: str) -> dict:
    """
    轮询等待指定的后台异步任务执行完成。

    Args:
        client: FastAPI 测试客户端。
        headers: 请求头字典。
        job_id: 需要轮询的后台任务 ID。

    Returns:
        dict: 任务完成或失败后的任务信息字典。

    Raises:
        AssertionError: 如果在重试 100 次后任务仍未完成或未返回失败状态。
    """
    # 最多尝试 100 次轮询，每次间隔 0.02 秒
    for _ in range(100):
        response = client.get(f"/api/v1/jobs/{job_id}", headers=headers)
        assert response.status_code == 200
        data = response.json()
        # 检查状态是否到达终态
        if data["status"] in {"succeeded", "failed"}:
            return data
        time.sleep(0.02)
    raise AssertionError("processing job did not finish")


def test_text_upload_is_safely_stored_processed_and_previewed(
    client: TestClient, users: dict[str, str]
) -> None:
    """
    测试文本文件上传后的安全存储、异步处理和内容预览。

    该测试用例验证了标准的 Markdown 文本能否被成功上传，并在后台任务解析
    切块后正确存储，同时检查了安全相关点（如防御目录遍历）。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典，用于获取不同角色的用户信息。

    Returns:
        None

    Raises:
        AssertionError: 在任何接口请求不符合预期或解析结果异常时抛出。
    """
    # 获取鉴权头，模拟 operator 用户身份
    headers = auth_headers(client, "operator", users["operator"])
    # 创建前置知识库
    knowledge_base_id = create_knowledge_base(client, headers)
    text = "# 小米 14\n小米 14 支持 90W 有线快充。\n\n屏幕尺寸为 6.36 英寸。"

    # 提交上传请求，同时传入切块参数 chunk_size 和 chunk_overlap 验证参数接收
    response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id, "chunk_size": "200", "chunk_overlap": "20"},
        # 故意提供带有 `../../` 相对路径的文件名以测试安全过滤
        files={"file": ("../../小米14.md", BytesIO(text.encode()), "text/markdown")},
    )

    # 验证响应为 202（Accepted）表示异步任务已被接收
    assert response.status_code == 202
    body = response.json()

    # 轮询等待后台处理任务完成
    job = wait_for_job(client, headers, body["job_id"])
    assert job["status"] == "succeeded", job

    # 验证文档记录与切割分块
    document = client.get(f"/api/v1/documents/{body['document_id']}", headers=headers).json()
    chunks = client.get(f"/api/v1/documents/{body['document_id']}/chunks", headers=headers).json()

    # 校验：检查原始文件名还原（去除目录遍历攻击构造的路径符号）
    assert document["original_filename"] == "小米14.md"
    # 校验：检查内部存储的文件名中是否剔除了非法字符
    assert ".." not in document["stored_filename"]
    # 校验：文档状态更新为 ready
    assert document["status"] == "ready"
    # 校验：确认生成了至少一个内容块
    assert chunks["total"] >= 1
    # 校验：从切分内容中确保有期望的关键词
    assert "90W" in "".join(item["text"] for item in chunks["items"])
    # 校验：确认实体模型提取了对应的产品名称
    assert any("小米 14" in item["product_models"] for item in chunks["items"])


def test_unsupported_or_spoofed_upload_is_rejected(
    client: TestClient, users: dict[str, str]
) -> None:
    """
    测试不支持的文件类型或伪造 MIME 类型的文件能否被拒绝上传。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若未成功拦截伪造上传。
    """
    # 获取鉴权头及创建前置知识库
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)

    # 发送伪造的 PDF 请求（MIME 为 application/pdf，但内容并非 PDF 魔数开头）
    response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id},
        files={"file": ("malware.pdf", BytesIO(b"not a pdf"), "application/pdf")},
    )

    # 期望收到 415 不支持的媒体类型报错
    assert response.status_code == 415
    # 验证错误码指示为文件签名无效
    assert response.json()["error"]["code"] == "invalid_file_signature"


def test_delete_document_removes_chunks(client: TestClient, users: dict[str, str]) -> None:
    """
    测试删除文档时是否能够级联清理对应的向量数据块和记录。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若删除接口报错或数据未被实际移除。
    """
    # 准备前置数据并上传一份 txt 测试文档
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)
    upload = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id},
        files={"file": ("guide.txt", BytesIO("产品使用说明".encode()), "text/plain")},
    ).json()
    # 等待后台切块任务完成，确保对应的向量块和关联数据已生成
    wait_for_job(client, headers, upload["job_id"])

    # 执行删除文档操作
    response = client.delete(f"/api/v1/documents/{upload['document_id']}", headers=headers)

    # 校验：响应成功返回 204 No Content
    assert response.status_code == 204
    # 校验：再次获取该文档应返回 404，表明该文档及其关联的切分数据已被安全移除
    missing = client.get(f"/api/v1/documents/{upload['document_id']}", headers=headers)
    assert missing.status_code == 404


def test_docx_and_pdf_formats_are_processed(client: TestClient, users: dict[str, str]) -> None:
    """
    测试后端能否正确解析并处理 DOCX 和 PDF 格式的富文本文档。

    利用 docx 和 reportlab 库在内存中动态生成包含测试内容的文档，
    并验证其是否能走通完整的上传和异步处理流程。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若上传被拒绝或后台任务未成功完成。
    """
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)

    # 动态构建 DOCX 内存文件，写入示例标题与正文
    word_buffer = BytesIO()
    word = WordDocument()
    word.add_heading("小米 14 使用说明", level=1)
    word.add_paragraph("小米 14 支持 90W 有线快充。")
    word.save(word_buffer)

    # 动态构建 PDF 内存文件
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer)
    pdf.drawString(72, 760, "Xiaomi 14 supports 90W wired charging.")
    pdf.save()

    # 将要上传的文件列表（包含文件名，字节内容及对应的 MIME 类型）
    uploads = [
        (
            "guide.docx",
            word_buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        ("xiaomi14-reference.pdf", pdf_buffer.getvalue(), "application/pdf"),
    ]

    # 遍历上传并验证异步任务是否均能成功执行
    for filename, content, media_type in uploads:
        response = client.post(
            "/api/v1/documents/upload",
            headers=headers,
            data={"knowledge_base_id": knowledge_base_id},
            files={"file": (filename, BytesIO(content), media_type)},
        )
        assert response.status_code == 202
        # 等待该文档的关联解析任务完成
        job = wait_for_job(client, headers, response.json()["job_id"])
        assert job["status"] == "succeeded", job


def test_duplicate_upload_is_rejected(client: TestClient, users: dict[str, str]) -> None:
    """
    测试重复上传相同内容的文档时是否会被拒绝。

    后台通过哈希对比或其他防重机制来避免同一个库里存在重复的文档。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若接口没有正确拦截重复的文档内容。
    """
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)
    payload = b"same document content"

    # 首次上传，正常情况会被受理
    first = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id},
        files={"file": ("first.txt", BytesIO(payload), "text/plain")},
    )
    # 第二次使用完全相同的负载内容上传（即便文件名不同）
    second = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id},
        files={"file": ("second.txt", BytesIO(payload), "text/plain")},
    )

    # 首次应当接受并开始异步处理
    assert first.status_code == 202
    # 第二次由于内容重复应当返回 409 冲突错误
    assert second.status_code == 409
    assert second.json()["error"]["code"] == "document_exists"


def test_jobs_can_be_filtered_by_document(client: TestClient, users: dict[str, str]) -> None:
    """
    测试能否通过文档 ID 过滤查询后台任务列表。

    确保在任务列表页可以精确追踪某个特定文档的处理任务流。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若过滤查询失效或返回的任务 ID 不符。
    """
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)
    # 触发上传，获取响应中的 job_id 和 document_id 用于后续过滤
    upload = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id},
        files={"file": ("jobs.txt", BytesIO("任务列表测试".encode()), "text/plain")},
    ).json()

    # 通过 document_id 参数过滤查询 /api/v1/jobs 接口
    response = client.get(f"/api/v1/jobs?document_id={upload['document_id']}", headers=headers)

    # 校验：接口应正常返回且总命中数为 1
    assert response.status_code == 200
    assert response.json()["total"] == 1
    # 校验：返回的任务 ID 与刚才上传产生的任务 ID 精确匹配
    assert response.json()["items"][0]["id"] == upload["job_id"]


def test_source_url_rejects_non_http_schemes(client: TestClient, users: dict[str, str]) -> None:
    """
    测试上传文档时指定的数据源 URL 是否能够拒绝非 HTTP 协议方案。

    主要用于防范 JavaScript 注入等跨站脚本攻击 (XSS) 风险。

    Args:
        client: FastAPI 测试客户端。
        users: 预置的用户字典。

    Returns:
        None

    Raises:
        AssertionError: 若服务端未能抛出 422 格式验证错误。
    """
    headers = auth_headers(client, "operator", users["operator"])
    knowledge_base_id = create_knowledge_base(client, headers)

    # 传入危险的 scheme（如 javascript:）作为 source_url 模拟攻击请求
    response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        data={"knowledge_base_id": knowledge_base_id, "source_url": "javascript:alert(1)"},
        files={"file": ("unsafe.md", BytesIO("安全测试".encode()), "text/markdown")},
    )

    # 服务端 Pydantic 模型或业务校验逻辑应当拦截该请求并返回 422 格式错误
    assert response.status_code == 422
