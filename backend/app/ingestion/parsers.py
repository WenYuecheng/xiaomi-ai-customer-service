"""
文件职责：
负责不同文件格式（pdf, docx, txt, md）的内容抽取和智能分块。

所属功能：
文档接入与处理 -> 解析与分块 (Chunking)。

主要流程：
1. 根据文件类型调用对应的提取库获取全文文本
2. 使用正则表达式和固定关键词进行产品实体识别
3. 使用 LangChain 文本切分器按特定分隔符切块
"""

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.ingestion.knowledge_quality import parse_front_matter


@dataclass(frozen=True)
class ParsedSection:
    """代表从原始文档中粗略抽取的片段及其物理位置描述（如页码）"""

    text: str
    location: str


def clean_text(value: str) -> str:
    """
    内部辅助函数：清理抽取出的文本，去除不可见字符、多余的空格和空行。

    主要职责：
    提高最终向量嵌入质量以及 LLM 提示词的纯净度。

    Args:
        value: 原始未经处理的抽取文本。

    Returns:
        清理完成的文本字符串。
    """
    value = unicodedata.normalize("NFKC", value).replace("\x00", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def load_sections(path: Path) -> list[ParsedSection]:
    """
    根据文件后缀，调度对应库（pypdf, python-docx）抽取出全文文本并附带简单位置。

    主要职责：
    作为文件解析入口，将非结构化文件转换成结构化的 `ParsedSection` 列表。

    Args:
        path: 本地文件系统中的待解析文件路径。

    Returns:
        包含提取到的文本及对应段落/页码位置的列表。

    Raises:
        ValueError: 当遇到不支持的文件后缀时抛出。
    """
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return [ParsedSection(clean_text(path.read_text(encoding="utf-8-sig")), "全文")]
    if suffix == ".pdf":
        reader = PdfReader(path)
        return [
            ParsedSection(clean_text(page.extract_text() or ""), f"第 {index} 页")
            for index, page in enumerate(reader.pages, start=1)
            if clean_text(page.extract_text() or "")
        ]
    if suffix == ".docx":
        document = WordDocument(path)
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return [ParsedSection(clean_text("\n".join(blocks)), "全文")]
    raise ValueError(f"unsupported document suffix: {suffix}")


PRODUCT_MODEL_PATTERNS = (
    re.compile(
        r"(?:小米|Xiaomi)\s*(\d+[A-Za-z]?(?:\s*(?:Pro|Ultra|Max))?)",
        re.IGNORECASE,
    ),
    re.compile(r"(?:红米|Redmi)\s*([A-Za-z]*\s*\d+(?:\s*(?:Pro|Ultra|Max))?)", re.IGNORECASE),
    re.compile(r"Smart\s+Band\s*(\d+(?:\s*Pro)?)", re.IGNORECASE),
    re.compile(r"Robot\s+Vacuum\s*(\d+(?:\s*Pro)?)", re.IGNORECASE),
    re.compile(r"(?:米家|Mijia)\s*([A-Za-z]*\s*\d+(?:\s*(?:Pro|Ultra|Max))?)", re.IGNORECASE),
    re.compile(r"\b([TX]\d+(?:\s*Pro)?)\b", re.IGNORECASE),
)


def extract_product_models(text: str) -> list[str]:
    """
    内部辅助函数：使用正则表达式，从切分后的文本块中提取产品型号名称。

    主要职责：
    用于建立知识图谱以及在检索（RAG）后期提供硬过滤所需的实体标签，提高检索准确率。

    Args:
        text: 待扫描的文本切块。

    Returns:
        提取到的所有不同产品型号列表，且自动去重并按字典序排列。
    """
    models: set[str] = set()
    metadata, _body = parse_front_matter(text)
    if model := metadata.get("型号", "").strip():
        return [re.sub(r"\s+", " ", model).strip().lower()]
    labels = ("小米", "红米", "Smart Band", "Robot Vacuum", "米家", "")
    for pattern, label in zip(PRODUCT_MODEL_PATTERNS, labels, strict=True):
        for match in pattern.findall(text):
            suffix = re.sub(r"\s+", " ", match).strip()
            models.add(f"{label} {suffix}".strip().lower())
    return sorted(models)


def split_sections(
    sections: list[ParsedSection], chunk_size: int, chunk_overlap: int
) -> list[tuple[str, str, list[str]]]:
    """
    对粗略提取出的章节，按字数和重叠长度进行平滑切块。

    主要职责：
    使用 Langchain 的 `RecursiveCharacterTextSplitter` 对文本进行拆解，以满足
    向量数据库及模型上下文限制要求，并同步调用 `extract_product_models` 获取对应切块产品实体。

    Args:
        sections: 粗提取得到的 `ParsedSection` 结构列表。
        chunk_size: 切块的最大字符数。
        chunk_overlap: 两个连续切块间的最大重叠字符数。

    Returns:
        返回 `(文本内容, 来源位置, 包含的产品实体列表)` 元组化结构构成的列表。
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[
            "\n# ",
            "\n## ",
            "\n\n",
            "\n",
            "。",
            "；",
            "，",
            "",
        ],
    )
    chunks: list[tuple[str, str, list[str]]] = []
    for section in sections:
        metadata, answerable_text = parse_front_matter(section.text)
        declared_model = re.sub(r"\s+", " ", metadata.get("型号", "")).strip().lower()
        for text in splitter.split_text(answerable_text):
            cleaned = clean_text(text)
            if cleaned:
                product_models = (
                    [declared_model] if declared_model else extract_product_models(cleaned)
                )
                chunks.append((cleaned, section.location, product_models))
    return chunks
